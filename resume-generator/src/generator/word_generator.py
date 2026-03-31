"""
Word简历生成器
基于python-docx库生成Word格式简历
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base import DocumentGenerator
from ..data.models import Person
import logging
import os

logger = logging.getLogger(__name__)


class WordGenerator(DocumentGenerator):
    """Word简历生成器"""
    
    def _load_template(self):
        """加载Word模板"""
        # 可以从现有.docx文件加载模板
        self.template_doc = None
        if os.path.exists(self.template_path):
            try:
                self.template_doc = Document(self.template_path)
                logger.info(f"成功加载Word模板: {self.template_path}")
            except Exception as e:
                logger.warning(f"加载模板失败，使用空白文档: {e}")
    
    def generate(self, person: Person, output_path: str, **kwargs) -> bool:
        """
        生成Word简历
        
        Args:
            person: 人员信息
            output_path: 输出文件路径
            **kwargs: 额外参数
            
        Returns:
            bool: 生成是否成功
        """
        try:
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 创建或加载文档
            if self.template_doc:
                doc = Document(self.template_path)
            else:
                doc = Document()
            
            # 清空现有内容（如果是基于模板）
            if self.template_doc:
                doc._body.clear()
            
            # 生成内容
            self._create_content(doc, person)
            
            # 保存文档
            doc.save(output_path)
            
            logger.info(f"成功生成Word简历: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"生成Word简历失败: {e}")
            return False
    
    def _create_content(self, doc: Document, person: Person):
        """创建Word内容"""
        # 标题
        title = doc.add_heading(person.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 联系信息
        doc.add_paragraph()
        contact_para = doc.add_paragraph()
        contact_para.add_run('联系信息').bold = True
        
        doc.add_paragraph(f'邮箱: {person.email}')
        doc.add_paragraph(f'电话: {person.phone}')
        doc.add_paragraph(f'学历: {person.education}')
        
        # 工作经历
        doc.add_paragraph()
        work_exp_heading = doc.add_paragraph()
        work_exp_heading.add_run('工作经历').bold = True
        work_exp_content = person.work_experience or '暂无工作经历信息'
        doc.add_paragraph(work_exp_content)
        
        # 专业技能
        if person.skills:
            doc.add_paragraph()
            skills_heading = doc.add_paragraph()
            skills_heading.add_run('专业技能').bold = True
            skills_content = ', '.join(person.skills)
            doc.add_paragraph(skills_content)
        
        # 相关证书
        if person.certifications:
            doc.add_paragraph()
            certs_heading = doc.add_paragraph()
            certs_heading.add_run('相关证书').bold = True
            certs_content = ', '.join(person.certifications)
            doc.add_paragraph(certs_content)